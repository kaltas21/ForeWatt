"""
Section 3: Conclusions, References, and Appendix for ForeWatt Technical Report.

This module generates the closing sections of the technical report including
conclusions, academic references, and appendix materials.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils_docx import (
    add_heading,
    add_paragraph,
    set_document_defaults,
    add_reference_entry,
    add_appendix_heading
)


def generate_section(doc: Document) -> None:
    """
    Generate Section 3 (Conclusions), References, and Appendix.

    Args:
        doc: The Document object to add content to.
    """
    _add_conclusions(doc)
    _add_references(doc)
    _add_appendix(doc)


def _add_conclusions(doc: Document) -> None:
    """Add the Conclusions section to the document."""

    add_heading(doc, "3. Conclusions", level=1)

    # Summary of Achievements
    add_paragraph(
        doc,
        "This technical report has presented ForeWatt, a production-ready electricity market "
        "forecasting platform designed specifically for the Turkish electricity market. The "
        "platform represents a comprehensive solution that integrates advanced machine learning "
        "techniques with robust data engineering practices to deliver accurate and reliable "
        "forecasts for both electricity consumption and day-ahead prices. Through systematic "
        "experimentation and rigorous evaluation, the developed system has demonstrated "
        "significant improvements over baseline forecasting methods, establishing its viability "
        "for real-world deployment in energy market applications."
    )

    add_paragraph(
        doc,
        "The consumption forecasting model achieved a symmetric Mean Absolute Percentage Error "
        "of 1.95 percent on the test dataset, representing a 2.8-fold improvement over the "
        "seasonal naive baseline. This level of accuracy was attained through careful feature "
        "engineering that incorporated temporal patterns, meteorological variables, and calendar "
        "effects specific to the Turkish electricity demand profile. The price forecasting model "
        "achieved an sMAPE of 11.71 percent, demonstrating a 1.8-fold improvement over the "
        "corresponding baseline. Given the inherently volatile and complex nature of electricity "
        "price dynamics, this performance represents a substantial achievement in capturing the "
        "fundamental drivers of price formation in the Turkish market."
    )

    add_paragraph(
        doc,
        "The evaluation of deep learning architectures, including Temporal Fusion Transformers, "
        "N-HITS, and PatchTST, revealed important insights regarding the applicability of these "
        "methods to medium-scale energy forecasting problems. Despite their demonstrated success "
        "in large-scale forecasting competitions, these architectures exhibited significant "
        "overfitting tendencies when applied to the available historical data spanning "
        "approximately two years. The gradient boosting ensemble approach, specifically the "
        "CHybrid V14 architecture combining CatBoost and LightGBM with context-aware error "
        "correction, proved to be the most effective methodology for this application domain."
    )

    # Technical Contributions
    add_paragraph(
        doc,
        "Several technical contributions have been made through the development of this platform. "
        "The medallion data architecture, organized into bronze, silver, and gold layers, provides "
        "a robust and maintainable framework for data ingestion, cleaning, and feature engineering. "
        "This architecture ensures data quality and traceability while facilitating the integration "
        "of multiple heterogeneous data sources including EPIAS market data, meteorological "
        "observations, and macroeconomic indicators. The feature engineering pipeline combines "
        "domain-specific knowledge of electricity markets with automated feature generation "
        "techniques, resulting in comprehensive feature sets that capture both physical and "
        "economic drivers of consumption and price dynamics."
    )

    add_paragraph(
        doc,
        "The production deployment on Google Cloud Run, coupled with the React-based dashboard "
        "and Firestore integration, demonstrates the practical applicability of the developed "
        "system. The hourly forecast updates, real-time visualization capabilities, and AI-powered "
        "chatbot interface provide stakeholders with actionable insights for market participation "
        "decisions. The anomaly detection module based on Isolation Forest algorithms further "
        "enhances the platform's utility by identifying unusual patterns that may require "
        "additional attention from market analysts."
    )

    # Future Work
    add_paragraph(
        doc,
        "Several directions for future work have been identified through this research. The "
        "expansion of the platform to cover additional electricity markets, particularly those "
        "in neighboring countries with interconnected grid infrastructure, represents a natural "
        "extension of the current capabilities. The integration of real-time news sentiment "
        "analysis could enhance price forecasting accuracy during periods of market stress or "
        "significant policy announcements. The implementation of automated model retraining "
        "pipelines would ensure that forecast accuracy is maintained as market dynamics evolve "
        "over time. Furthermore, the development of probabilistic forecasting capabilities "
        "through conformal prediction methods would provide decision-makers with calibrated "
        "uncertainty estimates essential for risk management applications. Enhanced anomaly "
        "detection incorporating temporal dependencies and market regime identification "
        "represents another promising avenue for improving the platform's analytical capabilities."
    )


def _add_references(doc: Document) -> None:
    """Add the References section to the document."""

    add_heading(doc, "References", level=1)

    references = [
        # Deep Learning for Time Series
        "Lim, B., Arık, S. Ö., Loeff, N., & Pfister, T. (2021). Temporal Fusion Transformers "
        "for interpretable multi-horizon time series forecasting. International Journal of "
        "Forecasting, 37(4), 1748-1764.",

        "Challu, C., Olivares, K. G., Oreshkin, B. N., Ramirez, F. G., Canseco, M. M., & "
        "Dubrawski, A. (2023). N-HiTS: Neural hierarchical interpolation for time series "
        "forecasting. Proceedings of the AAAI Conference on Artificial Intelligence, 37(6), "
        "6989-6997.",

        "Nie, Y., Nguyen, N. H., Sinthong, P., & Kalagnanam, J. (2023). A time series is worth "
        "64 words: Long-term forecasting with transformers. International Conference on Learning "
        "Representations (ICLR).",

        "Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., Kaiser, "
        "Ł., & Polosukhin, I. (2017). Attention is all you need. Advances in Neural Information "
        "Processing Systems, 30.",

        # Gradient Boosting Methods
        "Prokhorenkova, L., Gusev, G., Vorobev, A., Dorogush, A. V., & Gulin, A. (2018). "
        "CatBoost: Unbiased boosting with categorical features. Advances in Neural Information "
        "Processing Systems, 31.",

        "Ke, G., Meng, Q., Finley, T., Wang, T., Chen, W., Ma, W., Ye, Q., & Liu, T. Y. (2017). "
        "LightGBM: A highly efficient gradient boosting decision tree. Advances in Neural "
        "Information Processing Systems, 30.",

        "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings "
        "of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, "
        "785-794.",

        # Energy Forecasting
        "Weron, R. (2014). Electricity price forecasting: A review of the state-of-the-art with "
        "a look into the future. International Journal of Forecasting, 30(4), 1030-1081.",

        "Hong, T., & Fan, S. (2016). Probabilistic electric load forecasting: A tutorial review. "
        "International Journal of Forecasting, 32(3), 914-938.",

        "Lago, J., Marcjasz, G., De Schutter, B., & Weron, R. (2021). Forecasting day-ahead "
        "electricity prices: A review of state-of-the-art algorithms, best practices and an "
        "open-access benchmark. Applied Energy, 293, 116983.",

        "Ziel, F., & Weron, R. (2018). Day-ahead electricity price forecasting with high-dimensional "
        "structures: Univariate vs. multivariate modeling frameworks. Energy Economics, 70, 396-420.",

        "Nowotarski, J., & Weron, R. (2018). Recent advances in electricity price forecasting: "
        "A review of probabilistic forecasting. Renewable and Sustainable Energy Reviews, 81, "
        "1548-1568.",

        # Anomaly Detection
        "Liu, F. T., Ting, K. M., & Zhou, Z. H. (2008). Isolation forest. Proceedings of the "
        "2008 Eighth IEEE International Conference on Data Mining, 413-422.",

        # Turkish Electricity Market
        "EPIAS. (2024). EPIAS Transparency Platform documentation. Energy Exchange Istanbul. "
        "https://seffaflik.epias.com.tr/",

        "Kucukali, S., & Baris, K. (2010). Turkey's short-term gross annual electricity demand "
        "forecast by fuzzy logic approach. Energy Policy, 38(5), 2438-2445.",

        # Data Sources
        "Open-Meteo. (2024). Open-Meteo Weather API documentation. https://open-meteo.com/",

        "Central Bank of the Republic of Turkey. (2024). Electronic Data Delivery System (EVDS). "
        "https://evds2.tcmb.gov.tr/",

        # Software and Frameworks
        "Ramírez, S., & FastAPI Contributors. (2024). FastAPI: Modern, fast web framework for "
        "building APIs with Python. https://fastapi.tiangolo.com/",

        "React Contributors. (2024). React: A JavaScript library for building user interfaces. "
        "https://react.dev/",

        "Google Cloud. (2024). Cloud Run documentation. https://cloud.google.com/run/docs",

        "Firebase. (2024). Firebase documentation. https://firebase.google.com/docs",

        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., ... & "
        "Duchesnay, É. (2011). Scikit-learn: Machine learning in Python. Journal of Machine "
        "Learning Research, 12, 2825-2830.",

        "McKinney, W. (2010). Data structures for statistical computing in Python. Proceedings "
        "of the 9th Python in Science Conference, 56-61.",

        # Conformal Prediction
        "Vovk, V., Gammerman, A., & Shafer, G. (2005). Algorithmic learning in a random world. "
        "Springer Science & Business Media.",

        "Romano, Y., Patterson, E., & Candès, E. J. (2019). Conformalized quantile regression. "
        "Advances in Neural Information Processing Systems, 32.",

        # eptr2 Library
        "Korkut, R. (2024). eptr2: Python client for EPIAS Transparency Platform. "
        "https://github.com/Rogerspy/eptr2",
    ]

    for i, ref in enumerate(references, 1):
        add_reference_entry(doc, f"[{i}] {ref}")


def _add_appendix(doc: Document) -> None:
    """Add the Appendix section to the document."""

    add_heading(doc, "Appendix", level=1)

    # Appendix A: Feature Lists
    add_appendix_heading(doc, "A. Feature Specifications")

    add_paragraph(
        doc,
        "Complete feature lists for both consumption and price forecasting models are maintained "
        "in the project repository. The consumption model utilizes 23 features organized into "
        "four categories: temporal lag features capturing historical consumption patterns at "
        "24-hour, 48-hour, and 168-hour intervals; rolling statistical features including mean "
        "and standard deviation computed over 24-hour windows; meteorological features derived "
        "from weather observations across 12 Turkish cities; and calendar features encoding "
        "hour-of-day, day-of-week, month, weekend indicators, and Turkish public holidays. The "
        "price model employs 33 features encompassing price lags and rolling statistics, market "
        "fundamental indicators such as thermal gap and renewable saturation, system operation "
        "metrics including load factor and reserve margin ratios, and specialized features for "
        "solar generation patterns and price volatility dynamics."
    )

    # Appendix B: Hyperparameter Configurations
    add_appendix_heading(doc, "B. Hyperparameter Configurations")

    add_paragraph(
        doc,
        "The CatBoost models were trained with the following key hyperparameters: learning rate "
        "of 0.05, maximum depth of 8, L2 regularization coefficient of 3.0, and 2000 iterations "
        "with early stopping patience of 100 rounds. The LightGBM component of the price ensemble "
        "utilized a learning rate of 0.03, maximum depth of 10, number of leaves set to 64, and "
        "feature fraction of 0.8 for regularization. Both models employed symmetric quantile "
        "loss functions for training the prediction interval bounds. The Isolation Forest anomaly "
        "detection model was configured with 200 estimators, contamination parameter of 0.05, "
        "and bootstrap sampling enabled. All hyperparameters were determined through grid search "
        "with 5-fold time series cross-validation on the training dataset."
    )

    # Appendix C: API Reference
    add_appendix_heading(doc, "C. API Endpoint Reference")

    add_paragraph(
        doc,
        "The ForeWatt API exposes several endpoints for forecast retrieval and system monitoring. "
        "The health check endpoint (GET /health) returns the current system status and version "
        "information. The forecast trigger endpoint (POST /forecast) initiates the forecast "
        "pipeline and is invoked hourly by Google Cloud Scheduler. The latest forecast endpoint "
        "(GET /forecast/latest) returns the most recent 24-hour forecasts for both consumption "
        "and price models. Separate endpoints are available for retrieving price forecasts "
        "(GET /forecast/price) and consumption forecasts (GET /forecast/consumption) individually. "
        "Real-time data streaming is supported through the Firestore-backed endpoints "
        "(GET /api/realtime/{model}), while historical data access is provided via Parquet-based "
        "endpoints (GET /api/historical/{model}). Anomaly detection results are accessible through "
        "dedicated endpoints (GET /api/anomaly/{model}) that return identified anomalies with "
        "associated confidence scores and feature contributions."
    )

    add_paragraph(
        doc,
        "All API responses follow a standardized JSON schema with fields for timestamp, model "
        "version, forecast values, and optional confidence intervals. Authentication is handled "
        "through API keys for production deployments, while the development environment supports "
        "unauthenticated access for testing purposes. Rate limiting is implemented at 1000 "
        "requests per hour per client to ensure fair resource allocation among users."
    )


if __name__ == "__main__":
    # Create a standalone document for testing
    doc = Document()
    set_document_defaults(doc)

    generate_section(doc)

    output_path = "section_closing_test.docx"
    doc.save(output_path)
    print(f"Test document saved to {output_path}")
