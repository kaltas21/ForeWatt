"""
Shared utilities for ForeWatt Technical Report document generation.
Provides consistent formatting for all report sections using python-docx.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Report directory path
REPORT_DIR = os.path.dirname(os.path.abspath(__file__))


def create_document():
    """
    Create a new document with Times New Roman font and 1.5 line spacing.

    Returns:
        Document: A new python-docx Document with base styling configured.
    """
    doc = Document()

    # Set default font for Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Set paragraph formatting
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.space_after = Pt(6)

    # Configure heading styles
    for level in range(1, 4):
        style_name = f'Heading {level}'
        if style_name in doc.styles:
            heading_style = doc.styles[style_name]
            heading_font = heading_style.font
            heading_font.name = 'Times New Roman'
            heading_font.bold = True
            if level == 1:
                heading_font.size = Pt(14)
            elif level == 2:
                heading_font.size = Pt(13)
            else:
                heading_font.size = Pt(12)

    return doc


def set_font_recursive(element, font_name='Times New Roman'):
    """Set font for all runs in an element recursively."""
    for run in element.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_heading(doc, text, level=1):
    """
    Add a heading with proper Times New Roman formatting.

    Args:
        doc: Document object
        text: Heading text
        level: Heading level (1, 2, or 3)

    Returns:
        Paragraph: The heading paragraph
    """
    heading = doc.add_heading(text, level=level)

    # Set font for all runs in heading
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(14)
        elif level == 2:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(12)

    return heading


def add_paragraph(doc, text, bold=False, italic=False, first_line_indent=True):
    """
    Add a paragraph with standard styling.

    Args:
        doc: Document object
        text: Paragraph text
        bold: Whether text should be bold
        italic: Whether text should be italic
        first_line_indent: Whether to add first line indent

    Returns:
        Paragraph: The added paragraph
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic

    # Set paragraph formatting
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if first_line_indent:
        paragraph_format.first_line_indent = Cm(1.27)  # 0.5 inch

    return paragraph


def add_paragraph_with_runs(doc, runs_data, first_line_indent=True):
    """
    Add a paragraph with multiple runs (mixed formatting).

    Args:
        doc: Document object
        runs_data: List of tuples (text, bold, italic)
        first_line_indent: Whether to add first line indent

    Returns:
        Paragraph: The added paragraph
    """
    paragraph = doc.add_paragraph()

    for text, bold, italic in runs_data:
        run = paragraph.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic

    # Set paragraph formatting
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if first_line_indent:
        paragraph_format.first_line_indent = Cm(1.27)

    return paragraph


def add_figure(doc, image_path, caption, width_inches=5.5):
    """
    Add a figure with caption.

    Args:
        doc: Document object
        image_path: Path to image file (relative to Report directory)
        caption: Figure caption text
        width_inches: Width of image in inches

    Returns:
        tuple: (image_paragraph, caption_paragraph)
    """
    # Construct full path
    if not os.path.isabs(image_path):
        full_path = os.path.join(REPORT_DIR, image_path)
    else:
        full_path = image_path

    # Add image
    if os.path.exists(full_path):
        img_paragraph = doc.add_paragraph()
        img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_paragraph.add_run()
        run.add_picture(full_path, width=Inches(width_inches))
    else:
        img_paragraph = doc.add_paragraph(f"[Image not found: {image_path}]")
        img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add caption
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True

    # Add space after figure
    caption_paragraph.paragraph_format.space_after = Pt(12)

    return img_paragraph, caption_paragraph


def add_table(doc, data, headers, caption=None):
    """
    Add a formatted table with headers.

    Args:
        doc: Document object
        data: List of rows (each row is a list of cell values)
        headers: List of header strings
        caption: Optional table caption

    Returns:
        Table: The created table
    """
    # Add caption if provided
    if caption:
        caption_para = doc.add_paragraph()
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption_para.add_run(caption)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.bold = True
        caption_para.paragraph_format.space_after = Pt(6)

    # Create table
    num_cols = len(headers)
    num_rows = len(data) + 1  # +1 for header row

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Add headers
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = str(header)
        # Format header cell
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.bold = True
        # Set background color for header
        _set_cell_shading(cell, "D9D9D9")

    # Add data rows
    for row_idx, row_data in enumerate(data):
        row = table.rows[row_idx + 1]
        for col_idx, cell_value in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_value)
            # Format data cell
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)

    # Add space after table
    doc.add_paragraph()

    return table


def _set_cell_shading(cell, color):
    """Set cell background shading color."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


def add_title_page(doc, title, subtitle=None, authors=None, date=None):
    """
    Add a title page.

    Args:
        doc: Document object
        title: Main title
        subtitle: Optional subtitle
        authors: Optional list of author names
        date: Optional date string
    """
    # Add some space at top
    for _ in range(3):
        doc.add_paragraph()

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(24)
    run.bold = True

    # Subtitle
    if subtitle:
        doc.add_paragraph()
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_para.add_run(subtitle)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.italic = True

    # Add space before authors
    for _ in range(3):
        doc.add_paragraph()

    # Authors
    if authors:
        for author in authors:
            author_para = doc.add_paragraph()
            author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = author_para.add_run(author)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

    # Date
    if date:
        doc.add_paragraph()
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date_para.add_run(date)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


def add_abstract(doc, text):
    """
    Add an abstract section.

    Args:
        doc: Document object
        text: Abstract text
    """
    # Add heading
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("Abstract")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    # Add abstract text
    abstract_para = doc.add_paragraph()
    abstract_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = abstract_para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

    # Set paragraph formatting
    paragraph_format = abstract_para.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.left_indent = Cm(1.27)
    paragraph_format.right_indent = Cm(1.27)


def add_code_block(doc, code, language=None):
    """
    Add a code block.

    Args:
        doc: Document object
        code: Code text
        language: Optional language name for caption
    """
    code_para = doc.add_paragraph()
    code_para.paragraph_format.left_indent = Cm(1.27)
    code_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    run = code_para.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

    # Add shading to simulate code block
    # Note: Full background shading requires more complex XML manipulation


def save_section(doc, filename):
    """
    Save document to file in Report directory.

    Args:
        doc: Document object
        filename: Filename (will be saved in Report directory)
    """
    filepath = os.path.join(REPORT_DIR, filename)
    doc.save(filepath)
    print(f"Saved: {filepath}")
    return filepath


def get_figure_path(folder, filename):
    """
    Get full path to a figure file.

    Args:
        folder: Subfolder name (e.g., 'DesignFigures', 'ExperimentPlots')
        filename: Image filename

    Returns:
        str: Full path to the image file
    """
    return os.path.join(REPORT_DIR, folder, filename)


# Figure counter for automatic numbering
class FigureCounter:
    """Counter for figure numbering across sections."""

    def __init__(self, start=1):
        self.count = start

    def next(self):
        """Get next figure number and increment counter."""
        current = self.count
        self.count += 1
        return current

    def current(self):
        """Get current figure number without incrementing."""
        return self.count


# Table counter for automatic numbering
class TableCounter:
    """Counter for table numbering across sections."""

    def __init__(self, start=1):
        self.count = start

    def next(self):
        """Get next table number and increment counter."""
        current = self.count
        self.count += 1
        return current

    def current(self):
        """Get current table number without incrementing."""
        return self.count


def set_document_defaults(doc):
    """
    Set default document styling (alias for create_document styling on existing doc).

    Args:
        doc: Document object to configure
    """
    # Set default font for Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Set paragraph formatting
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.space_after = Pt(6)

    # Configure heading styles
    for level in range(1, 4):
        style_name = f'Heading {level}'
        if style_name in doc.styles:
            heading_style = doc.styles[style_name]
            heading_font = heading_style.font
            heading_font.name = 'Times New Roman'
            heading_font.bold = True
            if level == 1:
                heading_font.size = Pt(14)
            elif level == 2:
                heading_font.size = Pt(13)
            else:
                heading_font.size = Pt(12)


def add_reference_entry(doc, text):
    """
    Add a reference entry with hanging indent.

    Args:
        doc: Document object
        text: Reference text (e.g., "[1] Author, Title...")

    Returns:
        Paragraph: The reference paragraph
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

    # Set paragraph formatting with hanging indent
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.space_after = Pt(6)
    paragraph_format.left_indent = Cm(1.27)
    paragraph_format.first_line_indent = Cm(-1.27)  # Hanging indent

    return paragraph


def add_appendix_heading(doc, text):
    """
    Add an appendix subsection heading.

    Args:
        doc: Document object
        text: Heading text (e.g., "A. Feature Specifications")

    Returns:
        Paragraph: The heading paragraph
    """
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True

    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(12)
    paragraph_format.space_after = Pt(6)

    return paragraph


if __name__ == "__main__":
    # Test the utilities
    doc = create_document()
    add_title_page(doc, "ForeWatt Technical Report",
                   subtitle="Electricity Market Forecasting Platform",
                   authors=["ForeWatt Team"],
                   date="January 2026")
    add_page_break(doc)
    add_abstract(doc, "This is a test abstract for the ForeWatt Technical Report.")
    add_page_break(doc)
    add_heading(doc, "1. Introduction", level=1)
    add_paragraph(doc, "This is a test paragraph with proper formatting.")

    # Test table
    headers = ["Model", "sMAPE", "MAE", "R²"]
    data = [
        ["Consumption", "1.95%", "808.5 MWh", "0.969"],
        ["Price", "11.71%", "48.2 TL/MWh", "0.871"]
    ]
    add_table(doc, data, headers, caption="Table 1. Model Performance Summary")

    save_section(doc, "test_document.docx")
    print("Test completed successfully!")
